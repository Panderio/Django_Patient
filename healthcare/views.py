from django.core.mail import send_mail
from django.core.paginator import Paginator
from django.db.models.query import QuerySet
from django.shortcuts import get_object_or_404, redirect, render, reverse 
from django.contrib.auth.mixins import LoginRequiredMixin
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import authentication, permissions, generics, filters
from django.views import generic
from django.http import HttpResponse ,StreamingHttpResponse
from .models import Expert, Patient
from .forms import PatientModelForm , CustomUserCreation , PatientSearch , get_charts
import pandas as pd
from django.db.models import Q
import io
from docx import Document
from docx import *
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from docx.shared import Inches, Pt
from healthcare.api.serializers import PatientModelSeria
#from .worddoc import create_header



# Create your views here.

class SearchPatient(generics.ListAPIView):
    #Serializers needs Authentication
    serializer_class=PatientModelSeria
    def get_queryset(self):
        nom = self.request.query_params.get('first_name', None)
        print(nom)
        return Patient.objects.filter(first_name=nom)

class SearchPatientFilter(generics.ListAPIView):
    template_name="healthcare/patient_search.html"
    serializer_class=PatientModelSeria
    queryset=Patient.objects.all()
    context_object_name="patients"
    filter_backends=[filters.SearchFilter]
    search_fields=[
        '^Bureau_CNAM',
        '^médecin_conseil',
        '^gouvernorat',
        '^date_demande']
#Class Based Views
class ChartData(APIView,generic.TemplateView):
    authentication_classes = []
    permission_classes = []
    template_name="healthcare/patient_charts.html"
    def get(self,request,*args, **kwargs):
        qs = Patient.objects.all()
        serializer = PatientModelSeria(qs,many=True)
        return Response(serializer.data)


def charts(request):
        qs=Patient.objects.all()
        labels = []
        default_items = []

        for item in qs:
            labels.append(item.first_name)
            default_items.append(item.last_name)
            context = {
                "labels": labels,
                "default": default_items,
        }
        return render(request, "healthcare/patient_charts.html",context)

#USER
class SignupView(generic.CreateView):
    template_name="registration/signup.html"
    form_class=CustomUserCreation

    def get_success_url(self):
        return reverse("login")


#1
class LandingPageView(generic.TemplateView):
    template_name="landing.html"


#2
class PatientListView(LoginRequiredMixin, generic.ListView):
    template_name="healthcare/patient_list.html"
    context_object_name="patients"
    queryset=Patient.objects.all()
    paginate_by=8
    
    
#    def get_queryset(self):
#        user=self.request.user
#        if user.is_Expert:
#            queryset=Patient.objects.filter(docteur=user.userprofile)
#       else:
#          queryset=Patient.objects.filter(docteur=user.expert.docteur)
#         #filter for the conseil
#        queryset=queryset.filter(expert__user=user)
#   return queryset


#3
class PatientDetailView(LoginRequiredMixin, generic.DetailView):
    template_name="healthcare/patient_details.html"
    queryset=Patient.objects.all()
    context_object_name="patient"

#4
class PatientCreateView(LoginRequiredMixin, generic.CreateView):
    template_name="healthcare/patient_create.html"
    form_class=PatientModelForm

    def get_success_url(self):
        return reverse("healthcare:healthcare-patient_list")
    def form_valid(self,form):
        #TODO SENT EMAIL 
        send_mail(
            subject="A Patient Has Been Added", 
            message="check the List Patients for the new patient",
            from_email="test@test.com",
            recipient_list=["test2@test.com"] 
        )
        return super(PatientCreateView,self).form_valid(form)

#5
class PatientUpdateView(LoginRequiredMixin, generic.UpdateView):
    template_name="healthcare/patient_update.html"
    queryset=Patient.objects.all()
    form_class=PatientModelForm
    def get_success_url(self):
        return reverse("healthcare:healthcare-patient_list")


#6
class PatientDeleteView(LoginRequiredMixin, generic.DeleteView):
    template_name="healthcare/patient_delete.html"
    queryset=Patient.objects.all()
    context_object_name="patient"
    def get_success_url(self):
        return reverse("healthcare:healthcare-patient_list")


#Save Document !
class ExportDocx(PatientDetailView, APIView):
    def get(self, request,pk, *args, **kwargs):
        # create an empty document object
        document = Document()
        patient = Patient.objects.get(id=pk)
        print(patient)
        document = self.build_document(patient)

        # save document info
        buffer = io.BytesIO()
        document.save(buffer)  # save your memory stream
        buffer.seek(0)  # rewind the stream

        # put them to streaming content response 
        # within docx content_type
        response = StreamingHttpResponse(
            streaming_content=buffer,  # use the stream's content
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        response['Content-Disposition'] = 'attachment;filename=Download.docx'
        response["Content-Encoding"] = 'UTF-8'

        return response
    def build_document(self,queryset):
        print(queryset.civil)
        document = Document()
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(0.95)
            section.bottom_margin = Inches(0.95)
            section.left_margin = Inches(0.79)
            section.right_margin = Inches(0.79) 

        def create_header(name='Docteur Slah Eddine GHANNOUCHI',info='Professeur d’Anatomie à la Faculté de Médecine\nSpécialiste en Orthopédie\nCHU Farhat Hached\nExpert Judiciaire Assermenté près la Cour d’Appel Sousse\nMail:'):
            headerr = document.sections[0].header
            h_title=headerr.add_paragraph()
            h_title.paragraph_format.space_after = Pt(2)
            h_title=h_title.add_run(name)
            h2_title=headerr.add_paragraph().add_run(info)
            h_title.font.size = Pt(12)
            h2_title.font.size = Pt(8)
            h_title.font.name = 'Lucida Handwriting'
            h2_title.font.name = 'Lucida Handwriting'
            h_title.bold = True
            h2_title.bold = True
            h2_title.italic = True
            h_title.font.color.rgb = RGBColor(0,0,128)
            h2_title.font.color.rgb = RGBColor(0,0,128)
            return h_title,h2_title
        create_header()

        def dateNow():
            run = document.add_paragraph()
            run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            today = date.today()
            year = today.strftime("%Y")
            month = today.strftime("%m")
            day = today.strftime("%d")
            d = today.strftime("%d %B, %Y")
            todayy='Sousse le'+d
            datee= run.add_run(todayy)
            datee.font.size = Pt(14)
            datee.font.name = 'Times New Roman'
            return datee
        dateNow()

        def emptyparag():
            trashh=document.add_paragraph()
            trash=trashh.add_run('')
            trash.font.size = Pt(14)
            trash.font.name = 'Times New Roman'
            return trash

        def expertise():
            Header = document.add_paragraph()
            Header.alignment=WD_ALIGN_PARAGRAPH.CENTER
            Header = Header.add_run('EXPERTISE  MEDICALE DEMO')
            Header.font.size = Pt(26)
            Header.font.name = 'Times New Roman'
            Header.bold = True
            Header.font.color.rgb = RGBColor(0,0,0)
            return Header

        def firstParg():
            p = document.add_paragraph('')
            r = p.add_run('\tJe soussigné Docteur Slah-Eddine GHANNOUCHI, Professeur d’Anatomie à la Faculté de Médecine « IBN EL JAZZAR », Spécialiste en Orthopédie, exerçant au CHU Farhat HACHED de Sousse, Docteur en Biomécanique de l’Ecole Nationale Supérieure d’Arts et Métiers de Paris, titulaire du Diplôme d’Université de Réparation juridique du dommage corporel, de la Faculté de Médecine de Marseille, Expert judiciaire assermenté près la Cour d’Appel de Sousse,')
            r.font.size = Pt(14)
            r.font.name = 'Times New Roman'
            return r

        def SecndParag(M_F='Monsieur', docteur='',CNAM='',datee = date.today().strftime("%d %B, %Y")):
            p = document.add_paragraph('')
            str ='\tAgissant sur désignation de Madame le Docteur '+docteur+', Médecin Conseil de la Caisse Nationale d’Assurance Maladie à l’Unité du Contrôle Médical du Bureau de '+CNAM+', en date du'+datee
            r = p.add_run(str)
            r.font.size = Pt(14)
            r.font.name = 'Times New Roman'
            return r    
                # add a header
        document.add_heading(f'{queryset.first_name}\n{queryset.last_name}')

        # add a paragraph
        document.add_paragraph("This is a normal style paragraph")

        # add a paragraph within an italic text then go on with a break.
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.italic = True
        run.add_text("text will have italic style")
        run.add_break()
           
        return document

# Function Based Views
#1
def landing_page(request):
    return render(request, 'healthcare/landing.html')

def search_view(request):
    test = None
    data = None
    query_df= None
    chart=None
    form= PatientSearch(request.POST or None)
    if request.method =="POST":
        Bureau_CNAM = request.POST.get('Bureau_CNAM')
        médecin_conseil = request.POST.get('médecin_conseil')
        gouvernorat = request.POST.get('gouvernorat')
        date_demande = request.POST.get('date_demande')
        char_type = request.POST.get('char_type')
        print(Bureau_CNAM,médecin_conseil)
        data=[]
        queryset=Patient.objects.filter(médecin_conseil=médecin_conseil)
        if len(queryset)>0:
            query_df = pd.DataFrame(queryset.values())
            for qs in queryset:
                obj={
                    'id': qs.id,
                    'Nom': qs.last_name,
                    'prenom': qs.first_name,
                    'reprise': qs.reprise,
                    'durée_cotisation_an':qs.durée_cotisation_an

                }
                data.append(obj)
            data=pd.DataFrame(data)
            test=data.groupby('id', as_index=False)['durée_cotisation_an'].agg('sum')
            chart = get_charts(char_type,data,labels=data['durée_cotisation_an'].values)
            query_df = query_df.to_html()
            data = data.to_html()
            test = test.to_html()

        else:
            print('no data')
    context={
            'form':form,
            'query_df':query_df,
            'data':data,
            'test':test,
            'chart':chart
        }
    return render(request, 'healthcare/patient_search.html',context)
#2
def patient_searchh(request):
    search_q = request.GET.get('q','')
    if search_q:
        patients= Patient.objects.filter(Q(first_name=search_q) |
        Q(last_name=search_q)        
        )
    else:
        patients = Patient.objects.all()
    context={
        'patients':patients
    }
    return render(request, 'healthcare/about.html',context)

#3
def patient_details(request,pk):
    patient = Patient.objects.get(id=pk)
    context = {
        "patient":patient
    }
    return render(request, 'healthcare/patient_details.html',context)

#4
def patient_create(request):
    form=PatientModelForm()
    if request.method == "POST":
        form=PatientModelForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect("/")
    context = {
        "form":form
    }
    return render(request, 'healthcare/patient_create.html',context)


#5
def patient_update(request,pk):
    patient = Patient.objects.get(id=pk)
    form=PatientModelForm(instance=patient)
    if request.method == "POST":
        form=PatientModelForm(request.POST,instance=patient)
        if form.is_valid():
            form.save()
            return redirect("/")
    context = {
        "form":form,
        "patient":patient
    }
    return render(request, 'healthcare/patient_update.html',context)



def patient_delete(request,pk):
    patient = Patient.objects.get(id=pk)
    patient.delete()
    return redirect("/")








"""
OLD METHOD

def patient_create(request):
    form=PatientForm()
    if request.method == "POST":
        form=PatientForm(request.POST)
        if form.is_valid():
            first_name = form.cleaned_data["first_name"]
            last_name = form.cleaned_data["last_name"]
            civil = form.cleaned_data["civil"]
            matricule = form.cleaned_data["matricule"]
            cin = form.cleaned_data["cin"]
            age = form.cleaned_data["age"]
            Patient.objects.create(
                first_name=first_name,
                last_name=last_name,
                civil=civil,
                matricule=matricule,
                cin=cin,
                age=age
            )
            return redirect("/")
    context = {
        "form":form
    }
    return render(request, 'healthcare/patient_create.html',context)
"""