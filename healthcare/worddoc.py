from docx import *
document = Document()

def create_header(name='Docteur Slah Eddine GHANNOUCHI',info='Professeur d’Anatomie à la Faculté de Médecine\nSpécialiste en Orthopédie\nCHU Farhat Hached\nExpert Judiciaire Assermenté près la Cour d’Appel Sousse\nMail:'):
    headerr = document.sections[0].header
    h_title=headerr.add_paragraph()
    h_title.paragraph_format.space_after = Pt(2)
    h_title=h_title.add_run(name)
    h2_title=headerr.add_paragraph().add_run(info)
    h_title.font.size = Pt(12)
    h2_title.font.size = Pt(8)
    h_title.font.name = 'Lucida Handwriting'
    h2_title.font.name = 'Lucida Handwriting'
    h_title.bold = True
    h2_title.bold = True
    h2_title.italic = True
    h_title.font.color.rgb = RGBColor(0,0,128)
    h2_title.font.color.rgb = RGBColor(0,0,128)
    return h_title,h2_title

def dateNow():
    run = document.add_paragraph()
    run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    today = date.today()
    d = today.strftime("%d %B, %Y")
    todayy='Sousse le '+d
    datee= run.add_run(todayy)
    datee.font.size = Pt(14)
    datee.font.name = 'Times New Roman'
    return datee

def emptyparag():
    trashh=document.add_paragraph()
    trash=trashh.add_run('')
    trash.font.size = Pt(14)
    trash.font.name = 'Times New Roman'
    return trash

def expertise():
    Header = document.add_paragraph()
    Header.alignment=WD_ALIGN_PARAGRAPH.CENTER
    Header = Header.add_run('EXPERTISE  MEDICALE DEMO')
    Header.font.size = Pt(26)
    Header.font.name = 'Times New Roman'
    Header.bold = True
    Header.font.color.rgb = RGBColor(0,0,0)
    return Header


def firstParg():
    p = document.add_paragraph('')
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run('\tJe soussigné Docteur Slah-Eddine GHANNOUCHI, Professeur d’Anatomie à la Faculté de Médecine « IBN EL JAZZAR », Spécialiste en Orthopédie, exerçant au CHU Farhat HACHED de Sousse, Docteur en Biomécanique de l’Ecole Nationale Supérieure d’Arts et Métiers de Paris, titulaire du Diplôme d’Université de Réparation juridique du dommage corporel, de la Faculté de Médecine de Marseille, Expert judiciaire assermenté près la Cour d’Appel de Sousse,')
    r.font.size = Pt(14)
    r.font.name = 'Times New Roman'
    return r

def SecndParag(M_F='Monsieur', docteur='',CNAM='',datee = date.today().strftime("%d %B, %Y")):
    p = document.add_paragraph('')
    p.paragraph_format.space_after = Pt(0)
    str ='\tAgissant sur désignation de Madame le Docteur '+docteur+', Médecin Conseil de la Caisse Nationale d’Assurance Maladie à l’Unité du Contrôle Médical du Bureau de '+CNAM+', en date du'+datee
    r = p.add_run(str)
    r.font.size = Pt(14)
    r.font.name = 'Times New Roman'
    return r
#MARGIN PART 
sections = document.sections
for section in sections:
    section.top_margin = Inches(0.95)
    section.bottom_margin = Inches(0.95)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)

create_header('test','ahmed')
dateNow()
emptyparag()
emptyparag()
expertise()
firstParg()
SecndParag()
document.save('test.docx')